from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from typing import Dict, Any
from ...schemas.resume import ParsedResume, TransformRequest, TransformationResult
from ..deps import get_current_user
from ...services.ai_service import ai_service
import sys
import os

try:
    from multiverse_database.models import User
except ImportError:
    pass

router = APIRouter()

class GenerateVariantResponse(BaseModel):
    original_resume: Dict[str, Any]
    transformed_resume: Dict[str, Any]
    universe: str

@router.post("/generate_variant", response_model=GenerateVariantResponse)
async def generate_variant(
    universe: str = Form(...),
    file: UploadFile = File(...)
):
    # Temporarily remove auth dependency for easy frontend testing
    # current_user: "User" = Depends(get_current_user)
    
    if file.content_type not in ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "text/plain"]:
        raise HTTPException(400, "Unsupported file format. Please upload PDF, DOCX, or TXT.")
        
    content = await file.read()
    raw_text = ""
    
    if file.content_type == "application/pdf":
        try:
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            raw_text = "".join(page.extract_text() for page in reader.pages)
        except Exception as e:
            raise HTTPException(500, f"Error parsing PDF: {str(e)}")
            
    elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            from docx import Document
            from io import BytesIO
            doc = Document(BytesIO(content))
            raw_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as e:
            raise HTTPException(500, f"Error parsing DOCX: {str(e)}")
            
    else:
        raw_text = content.decode('utf-8')
        
    if not raw_text.strip():
        raise HTTPException(400, "Could not extract text from the file.")
        
    # 1. Parse Original Resume
    parsed_data = await ai_service.parse_resume(raw_text)
    
    # 2. Transform into Target Universe
    transformed_data = await ai_service.transform_resume(parsed_data, universe)
    
    return GenerateVariantResponse(
        original_resume=parsed_data,
        transformed_resume=transformed_data,
        universe=universe
    )

@router.post("/parse", response_model=ParsedResume)
async def parse_resume_upload(
    file: UploadFile = File(...),
    current_user: "User" = Depends(get_current_user)
):
    if file.content_type not in ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "text/plain"]:
        raise HTTPException(400, "Unsupported file format. Please upload PDF, DOCX, or TXT.")
        
    content = await file.read()
    raw_text = ""
    
    if file.content_type == "application/pdf":
        try:
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            raw_text = "".join(page.extract_text() for page in reader.pages)
        except Exception as e:
            raise HTTPException(500, f"Error parsing PDF: {str(e)}")
            
    elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            from docx import Document
            from io import BytesIO
            doc = Document(BytesIO(content))
            raw_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as e:
            raise HTTPException(500, f"Error parsing DOCX: {str(e)}")
            
    else:
        raw_text = content.decode('utf-8')
        
    if not raw_text.strip():
        raise HTTPException(400, "Could not extract text from the file.")
        
    parsed_data = await ai_service.parse_resume(raw_text)
    return ParsedResume(**parsed_data)

@router.post("/transform", response_model=TransformationResult)
async def transform_resume(
    request: TransformRequest,
    current_user: "User" = Depends(get_current_user)
):
    # In a real app, we would fetch the ParsedResume from DB using request.resume_id
    # For now, we mock the retrieval.
    mock_parsed_resume = {
        "personal_info": {"name": "John Doe", "email": "john@example.com", "phone": "1234567890"},
        "experience": [{"company": "Tech Corp", "role": "Software Engineer", "duration": "2020-2023", "description": "Built scalable backend systems."}],
        "education": [{"institution": "State University", "degree": "B.S. Computer Science", "year": "2020"}],
        "skills": ["Python", "FastAPI", "React", "PostgreSQL"],
        "projects": [{"name": "E-commerce platform", "description": "Developed a full-stack platform."}]
    }
    
    universe = request.universe_id # e.g., 'cyberpunk'
    
    transformed_data = await ai_service.transform_resume(mock_parsed_resume, universe)
    
    # Generate portrait
    resume_summary = f"{transformed_data.get('personal_info', {}).get('name', 'A professional')} working as {transformed_data.get('experience', [{}])[0].get('role', 'an expert')}"
    # Temporarily bypass actual image generation for cost/speed unless requested
    # portrait_url = await ai_service.generate_portrait(resume_summary, universe)
    portrait_url = "https://example.com/mock_portrait.png"
    
    return TransformationResult(
        original_resume_id=request.resume_id,
        universe_id=universe,
        transformed_content=transformed_data,
        portrait_url=portrait_url
    )
