from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from typing import Optional
import os
import uuid
from docx import Document
from ..deps import get_current_user

try:
    from multiverse_database.models import User
except ImportError:
    pass

router = APIRouter()

TEMP_DIR = "/tmp/multiverse_exports"
os.makedirs(TEMP_DIR, exist_ok=True)

@router.post("/docx/{generation_id}")
async def export_docx(
    generation_id: str,
    current_user: "User" = Depends(get_current_user)
):
    # Mock retrieving the generated resume
    # In reality, fetch from DB
    mock_content = {
        "personal_info": {"name": "Bruce Wayne (Batman)", "email": "bruce@wayne.enterprises"},
        "experience": [{"company": "Wayne Enterprises", "role": "CEO / Vigilante", "description": "Protecting Gotham."}]
    }
    
    filename = f"resume_{generation_id}.docx"
    filepath = os.path.join(TEMP_DIR, filename)
    
    doc = Document()
    doc.add_heading(mock_content["personal_info"]["name"], 0)
    doc.add_paragraph(mock_content["personal_info"]["email"])
    
    doc.add_heading('Experience', level=1)
    for exp in mock_content["experience"]:
        doc.add_heading(exp["company"], level=2)
        doc.add_heading(exp["role"], level=3)
        doc.add_paragraph(exp["description"])
        
    doc.save(filepath)
    
    return FileResponse(path=filepath, filename=filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@router.post("/pdf/{generation_id}")
async def export_pdf(
    generation_id: str,
    current_user: "User" = Depends(get_current_user)
):
    # Normally we would use Playwright to render an HTML template into a PDF
    # Mock implementation for scaffolding
    filename = f"resume_{generation_id}.pdf"
    filepath = os.path.join(TEMP_DIR, filename)
    
    # Placeholder: touch the file
    with open(filepath, 'w') as f:
        f.write("%PDF-1.4\n%Mock PDF Content")
        
    return FileResponse(path=filepath, filename=filename, media_type='application/pdf')
